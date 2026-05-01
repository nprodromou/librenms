from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()
BLACK = RGBColor(0, 0, 0)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLACK
    return h

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    return p

# Title block
title = doc.add_paragraph()
run = title.add_run("IFT 166: Introduction to Internet Networking")
run.bold = True
run.font.size = Pt(14)

subtitle = doc.add_paragraph()
run = subtitle.add_run("Lab 32 — Network Management with LibreNMS (Docker Alternative to ipMonitor / PRTG)")
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph(
    "This document describes an alternative path for Lab 32 for students who do not "
    "have access to a Windows machine. Instead of installing ipMonitor or PRTG, you "
    "will run LibreNMS in Docker. The learning outcome is the same: you will discover "
    "an SNMP-capable device and answer the same two questions about how SNMP communicates."
)

# Background
add_heading("What this lab demonstrates", level=1)
doc.add_paragraph(
    "SNMP (Simple Network Management Protocol) is how network administrators monitor "
    "switches, routers, servers, printers, and most other infrastructure. A monitoring "
    "tool sends an SNMP Get request to a device's SNMP agent, asking for a specific "
    "piece of information (uptime, interface counters, system description). The "
    "agent responds with the value."
)
doc.add_paragraph("In this lab you will:")
for item in [
    "Run a small Docker stack that includes LibreNMS (the network management system) and a simulated SNMP device.",
    "Discover the simulated device from LibreNMS, just as you would discover a real switch on a corporate network.",
    "Confirm the discovery worked by examining the device data in LibreNMS, then answer two questions about how SNMP communicates.",
]:
    doc.add_paragraph(item, style="List Number")

# What's running
add_heading("What you're installing", level=1)
doc.add_paragraph(
    "The compose.yml file in the repository brings up five containers:"
)
for item in [
    "librenms_db — MariaDB database that stores discovered devices and polled metrics.",
    "librenms_redis — Redis cache used by the LibreNMS web application.",
    "librenms — the LibreNMS web interface itself, served on http://localhost:8000.",
    "librenms_dispatcher — a worker process that performs the actual SNMP polling on a schedule.",
    "student — an SNMP simulator (snmpsim) that responds to queries as though it were a real Linux server. The container's name will reflect whatever you set as your student name in compose.yml.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Most home networks do not contain managed switches or other SNMP-enabled "
    "devices, so the simulator is included to guarantee that every student has a "
    "target to discover. The traffic LibreNMS sends to the simulator is identical to "
    "what it would send to a real Cisco switch on a corporate network."
)

# Setup
add_heading("Setup", level=1)

add_heading("Step 1. Install Docker Desktop", level=2)
doc.add_paragraph(
    "Download and install Docker Desktop from https://www.docker.com/products/docker-desktop. "
    "Follow the standard installation for your operating system. When the Docker icon "
    "in your menu bar (macOS) or system tray (Windows or Linux) shows that the engine "
    "is running, you are ready to continue."
)

add_heading("Step 2. Download the lab files", level=2)
doc.add_paragraph("If you already have git installed, clone the repository:")
add_code(
    "git clone https://github.com/nprodromou/librenms.git\n"
    "cd librenms\n"
    "cp .env.example .env"
)
doc.add_paragraph(
    "If you don't have git, download the tarball from "
    "https://github.com/nprodromou/librenms/archive/refs/heads/main.zip and "
    "double-click to extract it. Then in a terminal:"
)
add_code(
    "cd ~/Downloads/librenms-main\n"
    "cp .env.example .env"
)

add_heading("Step 3. Personalize your install", level=2)
doc.add_paragraph(
    "Open compose.yml in a text editor. At the very top of the file you will see:"
)
add_code(
    'x-vars:\n'
    '  student-name: &student_name "student"'
)
doc.add_paragraph(
    "Replace \"student\" with your name (lowercase, no spaces — e.g. \"jsmith\"). "
    "Save the file."
)
doc.add_paragraph("This single change personalizes three things:")
for item in [
    "The Docker container name (visible when you run docker compose ps).",
    "The hostname inside the Docker network (LibreNMS will discover it as a device with your name).",
    "The simulated device's sysName.0 SNMP response — your name appears as the device's system name in LibreNMS, which is how your instructor verifies the submission is yours.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "The third item is what allows your instructor to verify the submission is yours."
)

add_heading("Step 4. Bring up the stack", level=2)
add_code("docker compose up -d")
doc.add_paragraph(
    "The first run will download images and may take a few minutes. After the images "
    "are pulled, LibreNMS itself takes 60 to 90 seconds to initialize its database. "
    "You can watch progress with:"
)
add_code("docker compose logs -f librenms")
doc.add_paragraph(
    "When the log output settles, the web interface is ready."
)

add_heading("Step 5. Finish the LibreNMS install wizard", level=2)
doc.add_paragraph(
    "Open http://localhost:8000 in your browser. LibreNMS will run a one-time install "
    "wizard. Click through each step — the database connection should auto-fill from "
    "the environment variables. When asked, create an admin account using a username "
    "that identifies you (your ASU ID is a good choice). The username you create here "
    "is what your instructor will see in your screenshot, so use a real identifier."
)
doc.add_paragraph(
    "Insert a screenshot of the LibreNMS dashboard, with your username visible in the "
    "top-right corner of the page, below."
).italic = True
doc.add_paragraph("[INSERT SCREENSHOT — LibreNMS dashboard with your username visible]")

add_heading("Step 6. Add the simulated device", level=2)
doc.add_paragraph(
    "In the LibreNMS interface, navigate to Devices → Add Device. Fill in:"
)
for item in [
    "Hostname or IP: the name you set in the YAML anchor (e.g. jsmith).",
    "SNMP version: v2c.",
    "Community: public.",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "Click Add Device. LibreNMS will run discovery against the simulator. Within a "
    "few seconds you should see the device appear in the device list, populated with "
    "interfaces, system description, and other SNMP data."
)
doc.add_paragraph(
    "Insert a screenshot of the device list showing your personalized device, below."
).italic = True
doc.add_paragraph("[INSERT SCREENSHOT — LibreNMS device list with your personalized device]")

add_heading("Step 7. Confirm the SNMP discovery worked", level=2)
doc.add_paragraph(
    "Click on your device in the LibreNMS device list. The device detail page shows "
    "everything LibreNMS learned about it via SNMP: the system description, system "
    "name, contact, location, uptime, and a list of network interfaces with their "
    "current state and traffic counters."
)
doc.add_paragraph(
    "All of this information was retrieved by sending SNMP Get requests to the "
    "simulated device and parsing the responses. The fact that the page is populated "
    "is direct evidence that SNMP is working end-to-end between LibreNMS and the "
    "simulator."
)
doc.add_paragraph(
    "Insert a screenshot of the device detail page, below. Your personalized name "
    "should be visible as the device's system name."
).italic = True
doc.add_paragraph("[INSERT SCREENSHOT — LibreNMS device detail page showing SNMP-discovered data]")

# Lab questions
add_heading("Lab questions", level=1)
doc.add_paragraph(
    "Answer the following two questions in your own words. Use what you have read "
    "about SNMP and what you have observed working in LibreNMS to support your answer."
)

add_heading("Question 1: What source and destination port were used?", level=2)
doc.add_paragraph("Your answer:")
doc.add_paragraph("___________________________________________________________________")
doc.add_paragraph("___________________________________________________________________")

add_heading("Question 2: Why did the devices return Destination Unreachable ICMP packets?", level=2)
doc.add_paragraph("Your answer:")
doc.add_paragraph("___________________________________________________________________")
doc.add_paragraph("___________________________________________________________________")
doc.add_paragraph("___________________________________________________________________")

# Cleanup
add_heading("Stopping and cleaning up", level=1)
doc.add_paragraph("When you have completed the lab, you can stop the stack with:")
add_code("docker compose down")
doc.add_paragraph(
    "This stops all containers but keeps your data, so you can restart the stack "
    "later without re-running the install wizard. To completely remove the database "
    "and start fresh, run:"
)
add_code("docker compose down -v")

out = "/Users/nprodromou/Downloads/IFT166_Lab32_Prodromou_draft.docx"
doc.save(out)
print(f"Wrote {out}")
